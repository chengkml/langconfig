# Copyright (c) 2025 Cade Russell (Ghost Peony)
#
# This source code is licensed under the MIT license found in the
# LICENSE file in the root directory of this source tree.

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import PlainTextResponse, StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import desc
from sqlalchemy.exc import IntegrityError
from typing import List, Optional, Dict, Any
from pydantic import BaseModel, ConfigDict, Field, field_validator
from datetime import datetime, timedelta
import logging
from collections import defaultdict
import io
import re
import requests

from db.database import get_db
from models.workflow import WorkflowProfile, WorkflowStrategy, WorkflowVersion, WorkflowExecution
from models.execution_event import ExecutionEvent
from models.core import Task
from core.codegen.langgraph import LangGraphCodeGenerator
from core.session_manager import managed_transaction
from core.validation import ValidationError as CoreValidationError, PermissionError as CorePermissionError
from core.model_validators import workflow_validator
from core.versioning import check_version_conflict, format_lock_version_error

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/workflows", tags=["workflows"])


# Pydantic Schemas
class WorkflowProfileCreate(BaseModel):
    name: str
    description: Optional[str] = None
    project_id: Optional[int] = None  # Project association
    strategy_type: Optional[WorkflowStrategy] = None  # Optional - only needed for predefined strategy workflows
    configuration: dict = Field(...)  # Required - contains nodes and edges
    schema_output_config: Optional[dict] = None
    output_schema: Optional[str] = None
    blueprint: Optional[dict] = None
    custom_output_path: Optional[str] = None  # Custom output directory for file writes
    is_template: bool = False
    template_category: Optional[str] = None
    template_icon: Optional[str] = None
    template_tags: Optional[List[str]] = None


class WorkflowProfileUpdate(BaseModel):
    name: Optional[str] = None
    project_id: Optional[int] = None
    strategy_type: Optional[WorkflowStrategy] = None
    configuration: Optional[dict] = None
    schema_output_config: Optional[dict] = None
    output_schema: Optional[str] = None
    blueprint: Optional[dict] = None
    lock_version: Optional[int] = Field(None, description="Current lock version for optimistic locking")
    custom_output_path: Optional[str] = None  # Custom output directory for file writes
    is_template: Optional[bool] = None
    template_category: Optional[str] = None
    template_icon: Optional[str] = None
    template_tags: Optional[List[str]] = None


class WorkflowProfileResponse(BaseModel):
    model_config = ConfigDict(from_attributes=True)

    id: int
    name: str
    project_id: Optional[int]
    strategy_type: Optional[WorkflowStrategy]  # Optional - studio workflows don't need this
    configuration: dict
    schema_output_config: Optional[dict]
    output_schema: Optional[str]
    blueprint: Optional[dict]
    lock_version: int  # For optimistic locking
    custom_output_path: Optional[str]  # Custom output directory for file writes
    is_template: bool
    template_category: Optional[str]
    template_icon: Optional[str]
    template_tags: Optional[List[str]]
    created_at: datetime
    updated_at: datetime


class WorkflowForkRequest(BaseModel):
    name: Optional[str] = Field(None, description="Name for the forked workflow")
    project_id: Optional[int] = Field(None, description="Target project for the fork")
    as_template: bool = Field(False, description="Whether the fork should remain a reusable template")


class WorkflowTemplateRequest(BaseModel):
    is_template: bool = Field(True, description="Promote or demote this workflow as a template")
    category: Optional[str] = Field(None, description="Generic template category")
    icon: Optional[str] = Field(None, description="UI icon key")
    tags: Optional[List[str]] = Field(None, description="Searchable template tags")


# Endpoints
@router.get("/", response_model=List[WorkflowProfileResponse])
async def list_workflows(
    skip: int = 0,
    limit: int = 100,
    project_id: Optional[int] = None,
    is_template: Optional[bool] = None,
    template_category: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """List all workflow profiles, optionally filtered by project"""
    query = db.query(WorkflowProfile)

    if project_id is not None:
        query = query.filter(WorkflowProfile.project_id == project_id)
    if is_template is not None:
        query = query.filter(WorkflowProfile.is_template == is_template)
    if template_category:
        query = query.filter(WorkflowProfile.template_category == template_category)

    workflows = query.order_by(desc(WorkflowProfile.updated_at)).offset(skip).limit(limit).all()
    return workflows


@router.get("/{workflow_id}", response_model=WorkflowProfileResponse)
async def get_workflow(
    workflow_id: int,
    db: Session = Depends(get_db)
):
    """Get a specific workflow profile by ID"""
    workflow = db.query(WorkflowProfile).filter(WorkflowProfile.id == workflow_id).first()
    if not workflow:
        raise HTTPException(status_code=404, detail="Workflow not found")
    return workflow


@router.post("/", response_model=WorkflowProfileResponse, status_code=201)
async def create_workflow(
    workflow: WorkflowProfileCreate,
    db: Session = Depends(get_db)
):
    """
    Create a new workflow profile.

    Checks for duplicate names (case-insensitive) to prevent confusion.

    Raises:
        HTTPException 400: Workflow with this name already exists
        HTTPException 500: Server error
    """
    try:
        with managed_transaction(db, "create_workflow"):
            # QUICK WIN: Case-insensitive duplicate check
            # Prevents "MyWorkflow" and "myworkflow" as separate workflows
            existing = db.query(WorkflowProfile).filter(
                WorkflowProfile.name.ilike(workflow.name)
            ).first()

            if existing:
                raise HTTPException(
                    status_code=400,
                    detail=f"Workflow with name '{workflow.name}' already exists (case-insensitive match)"
                )

            db_workflow = WorkflowProfile(**workflow.model_dump())
            db.add(db_workflow)

        # Transaction commits here
        db.refresh(db_workflow)

        logger.info(
            f"Created workflow: {db_workflow.name}",
            extra={"workflow_id": db_workflow.id, "workflow_name": db_workflow.name}
        )

        return db_workflow

    except HTTPException:
        # Re-raise HTTP exceptions
        raise

    except Exception as e:
        # Unexpected error
        logger.error(f"Failed to create workflow: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create workflow: {str(e)}"
        )


@router.patch("/{workflow_id}", response_model=WorkflowProfileResponse)
async def update_workflow(
    workflow_id: int,
    workflow: WorkflowProfileUpdate,
    db: Session = Depends(get_db)
):
    """
    Update an existing workflow profile.

    Uses transaction safety to ensure workflow update succeeds even if auto-export fails.
    Export status is tracked in the database.

    Raises:
        HTTPException 404: Workflow not found
        HTTPException 500: Server error
    """
    try:
        with managed_transaction(db, f"update_workflow_{workflow_id}") as tx:
            db_workflow = db.query(WorkflowProfile).filter(
                WorkflowProfile.id == workflow_id
            ).first()

            if not db_workflow:
                raise HTTPException(status_code=404, detail="Workflow not found")

            # Optimistic locking - check lock version conflict
            if workflow.lock_version is not None:
                if check_version_conflict(db_workflow, workflow.lock_version):
                    error_msg = format_lock_version_error(
                        "Workflow",
                        workflow_id,
                        workflow.lock_version,
                        db_workflow.lock_version
                    )
                    raise HTTPException(status_code=409, detail=error_msg)

            # SMART AUTO-EXPORT: Capture old values before applying updates
            # We need to compare old vs new to detect actual changes
            old_blueprint = db_workflow.blueprint
            old_configuration = db_workflow.configuration

            # Update only provided fields
            update_data = workflow.model_dump(exclude_unset=True)

            # DEBUG: Log what we're receiving in configuration
            if "configuration" in update_data:
                config = update_data["configuration"]
                if isinstance(config, dict) and "nodes" in config:
                    for node in config.get("nodes", []):
                        node_id = node.get("id", "unknown")
                        node_config = node.get("config", {})
                        mcp_tools = node_config.get("mcp_tools", [])
                        cli_tools = node_config.get("cli_tools", [])
                        custom_tools = node_config.get("custom_tools", [])
                        logger.info(f"[PATCH /workflows/{workflow_id}] Node {node_id}:")
                        logger.info(f"  - mcp_tools: {mcp_tools}")
                        logger.info(f"  - cli_tools: {cli_tools}")
                        logger.info(f"  - custom_tools: {custom_tools}")
                        logger.info(f"  - config keys received: {list(node_config.keys())}")

            # CRITICAL FIX: Validate and filter updates using validation framework
            # This prevents unauthorized field updates (e.g., id, created_at, usage_count)
            try:
                validated_data = workflow_validator.validate_update(
                    update_data,
                    strict=True  # Reject unknown fields
                )
            except CoreValidationError as e:
                raise HTTPException(
                    status_code=422,
                    detail=f"Validation failed for field '{e.field}': {e.message}"
                )
            except CorePermissionError as e:
                raise HTTPException(
                    status_code=403,
                    detail=f"Permission denied for field '{e.field}': requires {e.permission.value} level"
                )

            # Apply validated updates (safe - only allowed fields)
            workflow_validator.apply_update(db_workflow, validated_data)

            # SMART AUTO-EXPORT: Only trigger export if blueprint/config ACTUALLY changed
            # Compare old vs new values to detect meaningful changes
            needs_export = _detect_export_triggering_changes(
                old_blueprint=old_blueprint,
                new_blueprint=db_workflow.blueprint,
                old_configuration=old_configuration,
                new_configuration=db_workflow.configuration,
                workflow=db_workflow
            )

            if needs_export:
                # Mark export as pending within transaction
                db_workflow.export_status = 'pending'
                db_workflow.export_error = None
                logger.info(f"Workflow {workflow_id} needs export, marking as pending")

        # Transaction commits here - workflow update and export status saved together
        db.refresh(db_workflow)

        # Trigger export in background using task queue
        if needs_export:
            try:
                from core.task_queue import task_queue, TaskPriority

                # Enqueue auto-export task (non-blocking)
                task_id = await task_queue.enqueue(
                    "auto_export_workflow",
                    {"workflow_id": workflow_id},
                    priority=TaskPriority.NORMAL,
                    max_retries=3
                )

                logger.info(
                    f"Auto-export task enqueued for workflow {workflow_id} (task_id: {task_id})",
                    extra={
                        "workflow_id": workflow_id,
                        "task_id": task_id
                    }
                )

            except Exception as e:
                # Failed to enqueue task - update export status
                logger.error(
                    f"Failed to enqueue auto-export for workflow {workflow_id}: {e}",
                    exc_info=True
                )
                # Update status to failed
                try:
                    db_workflow.export_status = 'failed'
                    db_workflow.export_error = f"Failed to enqueue export: {str(e)}"
                    db.commit()
                    db.refresh(db_workflow)
                except Exception as update_error:
                    logger.error(f"Failed to update export status: {update_error}")

        logger.info(
            f"Updated workflow {workflow_id}",
            extra={
                "workflow_id": workflow_id,
                "needs_export": needs_export,
                "export_status": db_workflow.export_status
            }
        )

        return db_workflow

    except HTTPException:
        # Re-raise HTTP exceptions
        raise

    except Exception as e:
        # Unexpected error
        logger.error(
            f"Failed to update workflow {workflow_id}: {e}",
            exc_info=True,
            extra={"workflow_id": workflow_id}
        )
        raise HTTPException(
            status_code=500,
            detail=f"Failed to update workflow: {str(e)}"
        )


def _escape_like_pattern(value: str) -> str:
    """Escape LIKE/ILIKE wildcards so user input matches literally."""
    return value.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")


def _unique_workflow_name(db: Session, requested_name: str) -> str:
    """Return a workflow name that does not collide with existing profiles."""
    base_name = requested_name.strip()
    if not base_name:
        base_name = "Untitled Workflow"

    escaped_base = _escape_like_pattern(base_name)
    existing = db.query(WorkflowProfile).filter(
        WorkflowProfile.name.ilike(escaped_base, escape="\\")
    ).first()
    if not existing:
        return base_name

    counter = 2
    while db.query(WorkflowProfile).filter(
        WorkflowProfile.name.ilike(f"{escaped_base} ({counter})", escape="\\")
    ).first():
        counter += 1
    return f"{base_name} ({counter})"


@router.post("/{workflow_id}/fork", response_model=WorkflowProfileResponse, status_code=201)
async def fork_workflow(
    workflow_id: int,
    request: WorkflowForkRequest,
    db: Session = Depends(get_db)
):
    """Create an independent editable copy of a workflow or template."""
    source = db.query(WorkflowProfile).filter(WorkflowProfile.id == workflow_id).first()
    if not source:
        raise HTTPException(status_code=404, detail="Source workflow not found")

    fork_name = _unique_workflow_name(db, request.name or f"Copy of {source.name}")
    fork = WorkflowProfile(
        name=fork_name,
        description=source.description,
        project_id=request.project_id,
        strategy_type=source.strategy_type,
        configuration=source.configuration or {"nodes": [], "edges": []},
        schema_output_config=source.schema_output_config,
        output_schema=source.output_schema,
        blueprint=source.blueprint,
        custom_output_path=source.custom_output_path,
        is_template=request.as_template,
        template_category=source.template_category if request.as_template else None,
        template_icon=source.template_icon if request.as_template else None,
        template_tags=source.template_tags if request.as_template else None,
    )
    db.add(fork)
    db.commit()
    db.refresh(fork)
    return fork


@router.patch("/{workflow_id}/template", response_model=WorkflowProfileResponse)
async def update_workflow_template_status(
    workflow_id: int,
    request: WorkflowTemplateRequest,
    db: Session = Depends(get_db)
):
    """Promote a workflow to a reusable template or demote it back to a regular workflow."""
    workflow = db.query(WorkflowProfile).filter(WorkflowProfile.id == workflow_id).first()
    if not workflow:
        raise HTTPException(status_code=404, detail="Workflow not found")

    workflow.is_template = request.is_template
    if request.is_template:
        if request.category is not None:
            workflow.template_category = request.category
        if request.icon is not None:
            workflow.template_icon = request.icon
        if request.tags is not None:
            workflow.template_tags = request.tags
    else:
        workflow.template_category = None
        workflow.template_icon = None
        workflow.template_tags = None

    db.commit()
    db.refresh(workflow)
    return workflow


def _should_auto_export(workflow: WorkflowProfile) -> bool:
    """
    Determine if workflow needs auto-export.

    Checks if workflow contains DeepAgent nodes that require export.
    DeepAgents are identified by:
    - node.type == 'deepagent'
    - node.data.use_deepagents == True
    - node.data.config.use_deepagents == True
    - node.data.config.middleware (has DeepAgent middleware configured)
    - node.data.deepAgentId (references a DeepAgent template)
    - node.data.agentType == 'deep' (agent type selector)

    Args:
        workflow: WorkflowProfile instance

    Returns:
        bool: True if workflow needs export, False otherwise
    """
    if not workflow.blueprint:
        return False

    # Check for DeepAgent nodes
    if isinstance(workflow.blueprint, dict):
        nodes = workflow.blueprint.get('nodes', [])
        logger.debug(f"Checking {len(nodes)} nodes for DeepAgent indicators")
        for node in nodes:
            node_id = node.get('id', 'unknown')
            node_type = node.get('type', 'unknown')
            node_data = node.get('data', {})
            logger.debug(f"  Node {node_id}: type={node_type}, data_keys={list(node_data.keys())}")
            # Check explicit type
            if node.get('type') == 'deepagent':
                logger.debug(f"Node {node.get('id')} is deepagent type")
                return True

            # Check use_deepagents flag in data
            node_data = node.get('data', {})
            if node_data.get('use_deepagents'):
                logger.debug(f"Node {node.get('id')} has use_deepagents=True in data")
                return True

            # Check if node references a DeepAgent template (deepAgentId)
            if node_data.get('deepAgentId') or node_data.get('deep_agent_id'):
                logger.debug(f"Node {node.get('id')} references DeepAgent template")
                return True

            # Check agentType field (from agent builder)
            agent_type = node_data.get('agentType') or node_data.get('agent_type')
            if agent_type == 'deep':
                logger.debug(f"Node {node.get('id')} has agentType=deep")
                return True

            # Check use_deepagents in config
            config = node_data.get('config', {})
            if config.get('use_deepagents'):
                logger.debug(f"Node {node.get('id')} has use_deepagents=True in config")
                return True

            # Check agentType in config
            config_agent_type = config.get('agentType') or config.get('agent_type')
            if config_agent_type == 'deep':
                logger.debug(f"Node {node.get('id')} has config.agentType=deep")
                return True

            # Check for DeepAgent middleware (todo_list, filesystem, subagent)
            middleware = config.get('middleware', [])
            if middleware:
                # If middleware is configured, it's likely a DeepAgent
                middleware_types = [m.get('type') for m in middleware if isinstance(m, dict)]
                deepagent_middleware = {'todo_list', 'filesystem', 'subagent'}
                if set(middleware_types) & deepagent_middleware:
                    logger.debug(f"Node {node.get('id')} has DeepAgent middleware: {middleware_types}")
                    return True

            # Check for subagents configuration
            if config.get('subagents') or node_data.get('subagents'):
                logger.debug(f"Node {node.get('id')} has subagents configured")
                return True

    return False


def _detect_export_triggering_changes(
    old_blueprint: Optional[dict],
    new_blueprint: Optional[dict],
    old_configuration: Optional[dict],
    new_configuration: Optional[dict],
    workflow: WorkflowProfile
) -> bool:
    """
    Detect if blueprint or configuration changes require workflow export.

    Smart export triggering:
    - Only triggers if workflow has DeepAgent nodes
    - Only triggers if blueprint or configuration ACTUALLY changed
    - Prevents unnecessary exports when updating metadata fields (name, description, etc.)

    Args:
        old_blueprint: Previous blueprint value
        new_blueprint: New blueprint value
        old_configuration: Previous configuration value
        new_configuration: New configuration value
        workflow: WorkflowProfile instance (for checking if export is needed)

    Returns:
        bool: True if export should be triggered, False otherwise
    """
    # First check: Does this workflow even need export? (has DeepAgent nodes)
    if not _should_auto_export(workflow):
        logger.debug(f"Workflow {workflow.id} does not have DeepAgent nodes, skipping export")
        return False

    # Second check: Did blueprint actually change?
    blueprint_changed = False
    if old_blueprint != new_blueprint:
        blueprint_changed = True
        logger.info(
            f"Workflow {workflow.id} blueprint changed, export needed",
            extra={
                "workflow_id": workflow.id,
                "old_blueprint_keys": list(old_blueprint.keys()) if old_blueprint else [],
                "new_blueprint_keys": list(new_blueprint.keys()) if new_blueprint else []
            }
        )

    # Third check: Did configuration actually change?
    configuration_changed = False
    if old_configuration != new_configuration:
        configuration_changed = True
        logger.info(
            f"Workflow {workflow.id} configuration changed, export needed",
            extra={
                "workflow_id": workflow.id,
                "old_config_keys": list(old_configuration.keys()) if old_configuration else [],
                "new_config_keys": list(new_configuration.keys()) if new_configuration else []
            }
        )

    # Trigger export only if something actually changed
    if blueprint_changed or configuration_changed:
        logger.info(
            f"Export triggered for workflow {workflow.id}",
            extra={
                "blueprint_changed": blueprint_changed,
                "configuration_changed": configuration_changed
            }
        )
        return True

    # No meaningful changes detected
    logger.debug(
        f"No export-triggering changes for workflow {workflow.id}",
        extra={"blueprint_changed": False, "configuration_changed": False}
    )
    return False


def cleanup_execution_history(
    db: Session,
    workflow_id: int,
    max_count: Optional[int] = None,
    retention_days: Optional[int] = None
) -> Dict[str, int]:
    """
    Clean up old execution history for a workflow.

    Removes executions that exceed configured limits:
    - Count-based: Keep only the N most recent executions
    - Age-based: Remove executions older than retention period

    Configurable execution history limits
    - Prevents unbounded growth of execution records
    - Improves query performance for execution history
    - Configurable via environment variables

    Args:
        db: Database session
        workflow_id: ID of the workflow to clean up
        max_count: Maximum executions to keep (defaults to settings value)
        retention_days: Max age in days (defaults to settings value)

    Returns:
        Dict with cleanup statistics:
        {
            "removed_by_count": int,
            "removed_by_age": int,
            "total_removed": int,
            "remaining": int
        }
    """
    from config import settings

    max_count = max_count or settings.max_execution_history_per_workflow
    retention_days = retention_days or settings.execution_history_retention_days

    removed_by_count = 0
    removed_by_age = 0

    try:
        # Cleanup 1: Remove executions beyond max count
        # Keep the N most recent executions, delete the rest
        if max_count > 0:
            # Get all executions for this workflow, ordered by completion time
            all_executions = db.query(WorkflowExecution).filter(
                WorkflowExecution.workflow_id == workflow_id
            ).order_by(desc(WorkflowExecution.completed_at)).all()

            # If we have more than max_count, delete the excess
            if len(all_executions) > max_count:
                executions_to_delete = all_executions[max_count:]
                execution_ids_to_delete = [e.id for e in executions_to_delete]

                # Delete in batch
                deleted = db.query(WorkflowExecution).filter(
                    WorkflowExecution.id.in_(execution_ids_to_delete)
                ).delete(synchronize_session=False)

                removed_by_count = deleted
                logger.info(
                    f"Removed {deleted} old executions for workflow {workflow_id} (count limit: {max_count})",
                    extra={
                        "workflow_id": workflow_id,
                        "removed_count": deleted,
                        "max_count": max_count
                    }
                )

        # Cleanup 2: Remove executions older than retention period
        if retention_days > 0:
            cutoff_date = datetime.utcnow() - timedelta(days=retention_days)

            deleted = db.query(WorkflowExecution).filter(
                WorkflowExecution.workflow_id == workflow_id,
                WorkflowExecution.completed_at < cutoff_date
            ).delete(synchronize_session=False)

            removed_by_age = deleted
            logger.info(
                f"Removed {deleted} expired executions for workflow {workflow_id} (older than {retention_days} days)",
                extra={
                    "workflow_id": workflow_id,
                    "removed_count": deleted,
                    "retention_days": retention_days,
                    "cutoff_date": cutoff_date.isoformat()
                }
            )

        # Commit cleanup changes
        db.commit()

        # Count remaining executions
        remaining = db.query(WorkflowExecution).filter(
            WorkflowExecution.workflow_id == workflow_id
        ).count()

        total_removed = removed_by_count + removed_by_age

        logger.info(
            f"Execution history cleanup completed for workflow {workflow_id}",
            extra={
                "workflow_id": workflow_id,
                "removed_by_count": removed_by_count,
                "removed_by_age": removed_by_age,
                "total_removed": total_removed,
                "remaining": remaining
            }
        )

        return {
            "removed_by_count": removed_by_count,
            "removed_by_age": removed_by_age,
            "total_removed": total_removed,
            "remaining": remaining
        }

    except Exception as e:
        db.rollback()
        logger.error(
            f"Failed to cleanup execution history for workflow {workflow_id}: {e}",
            exc_info=True,
            extra={"workflow_id": workflow_id}
        )
        # Return empty stats on error
        return {
            "removed_by_count": 0,
            "removed_by_age": 0,
            "total_removed": 0,
            "remaining": 0
        }


@router.delete("/{workflow_id}", status_code=204)
async def delete_workflow(
    workflow_id: int,
    db: Session = Depends(get_db)
):
    """
    Delete a workflow profile and all related data.

    Cascades to delete:
    - Tasks associated with the workflow
    - Execution events from those tasks
    - Generated files in the workspace

    Note: Memory is managed by LangGraph's store system (per-thread)
    and doesn't require explicit cleanup.
    """
    db_workflow = db.query(WorkflowProfile).filter(WorkflowProfile.id == workflow_id).first()
    if not db_workflow:
        raise HTTPException(status_code=404, detail="Workflow not found")

    try:
        # Delete related data in order of dependencies

        # 1. Delete execution events for all tasks in this workflow
        db.query(ExecutionEvent).filter(
            ExecutionEvent.workflow_id == workflow_id
        ).delete(synchronize_session=False)

        # 2. Delete tasks associated with this workflow
        # Use workflow_profile_id (the actual foreign key), not workflow_id (tracking string)
        db.query(Task).filter(
            Task.workflow_profile_id == workflow_id
        ).delete(synchronize_session=False)

        # 3. Memory is stored in LangGraph's store system (not a separate table)
        # It's managed per-thread and doesn't need explicit cleanup

        # 4. Delete workflow executions (must come before workflow versions)
        db.query(WorkflowExecution).filter(
            WorkflowExecution.workflow_id == workflow_id
        ).delete(synchronize_session=False)

        # 5. Delete workflow versions
        db.query(WorkflowVersion).filter(
            WorkflowVersion.workflow_id == workflow_id
        ).delete(synchronize_session=False)

        # 6. Finally, delete the workflow itself
        db.delete(db_workflow)

        db.commit()
        logger.info(f"Successfully deleted workflow {workflow_id} and all related data")
        return None

    except Exception as e:
        db.rollback()
        logger.error(f"Failed to delete workflow {workflow_id}: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to delete workflow: {str(e)}"
        )


@router.get("/{workflow_id}/code", response_class=PlainTextResponse)
async def generate_workflow_code(
    workflow_id: int,
    db: Session = Depends(get_db)
):
    """Generate standalone LangGraph Python code for a workflow"""
    # Get the workflow
    workflow = db.query(WorkflowProfile).filter(WorkflowProfile.id == workflow_id).first()
    if not workflow:
        raise HTTPException(status_code=404, detail="Workflow not found")

    # Convert workflow to dict
    workflow_dict = {
        'id': workflow.id,
        'name': workflow.name,
        'description': f'Auto-generated workflow from {workflow.name}',
        'strategy_type': workflow.strategy_type.value if hasattr(workflow.strategy_type, 'value') else str(workflow.strategy_type),
        'blueprint': workflow.blueprint or {},
        'configuration': workflow.configuration or {}
    }

    # Generate code
    generator = LangGraphCodeGenerator()
    generated_code = generator.generate_from_workflow(workflow_dict)

    return generated_code


# Helper function for auto-exporting DeepAgent workflows
async def auto_export_deepagent_workflow(workflow: WorkflowProfile, db: Session):
    """
    Automatically generate exports for workflows that contain DeepAgents.

    This checks if the workflow contains any DeepAgent configurations and
    generates both standalone and .langconfig exports automatically.
    """
    # Check if workflow has DeepAgent configurations in its blueprint
    if not workflow.blueprint:
        return

    # Look for DeepAgent nodes in the blueprint using same logic as _should_auto_export
    has_deepagent = False
    deepagent_nodes = []

    if isinstance(workflow.blueprint, dict):
        nodes = workflow.blueprint.get('nodes', [])
        for node in nodes:
            node_data = node.get('data', {})
            config = node_data.get('config', {})

            # Check all DeepAgent indicators
            is_deepagent = (
                node.get('type') == 'deepagent' or
                node_data.get('use_deepagents') or
                node_data.get('deepAgentId') or
                node_data.get('deep_agent_id') or
                node_data.get('agentType') == 'deep' or
                node_data.get('agent_type') == 'deep' or
                config.get('use_deepagents') or
                config.get('agentType') == 'deep' or
                config.get('agent_type') == 'deep' or
                config.get('subagents') or
                node_data.get('subagents')
            )

            # Also check middleware
            middleware = config.get('middleware', [])
            if middleware:
                middleware_types = [m.get('type') for m in middleware if isinstance(m, dict)]
                deepagent_middleware = {'todo_list', 'filesystem', 'subagent'}
                if set(middleware_types) & deepagent_middleware:
                    is_deepagent = True

            if is_deepagent:
                has_deepagent = True
                deepagent_nodes.append(node)

    if not has_deepagent:
        return

    logger.info(f"Auto-exporting DeepAgent workflow: {workflow.name}")

    # Import here to avoid circular dependencies
    from models.deep_agent import DeepAgentTemplate, DeepAgentConfig
    from services.export_service import ExportService

    # For each DeepAgent node, generate exports
    for node in deepagent_nodes:
        try:
            # Get agent configuration from node data
            agent_data = node.get('data', {})

            # Check if this agent is already in the database
            agent_name = agent_data.get('name', f"Workflow Agent - {workflow.name}")
            agent = db.query(DeepAgentTemplate).filter(
                DeepAgentTemplate.name == agent_name
            ).first()

            if agent:
                # Parse configuration
                config = DeepAgentConfig(**agent.config)

                # Generate both export formats
                try:
                    # Standalone export
                    await ExportService.export_standalone(agent, config)
                    logger.info(f"Generated standalone export for {agent.name}")
                except Exception as e:
                    logger.error(f"Standalone export failed: {e}")

                try:
                    # LangConfig export
                    await ExportService.export_langconfig(agent, config)
                    logger.info(f"Generated .langconfig export for {agent.name}")
                except Exception as e:
                    logger.error(f"LangConfig export failed: {e}")

        except Exception as e:
            logger.error(f"Error exporting DeepAgent node: {e}")
            continue

    logger.info(f"Auto-export complete for workflow: {workflow.name}")


# =============================================================================
# Output Path Configuration Endpoints
# =============================================================================

class ValidatePathRequest(BaseModel):
    path: str = Field(..., description="The custom output path to validate")


class ValidatePathResponse(BaseModel):
    valid: bool
    resolved_path: Optional[str] = None
    error: Optional[str] = None
    writable: bool = False
    exists: bool = False


# Blocked path patterns for security
BLOCKED_PATH_PATTERNS = [
    # Windows system directories
    r'^[A-Za-z]:\\Windows',
    r'^[A-Za-z]:\\Program Files',
    r'^[A-Za-z]:\\Program Files \(x86\)',
    # Unix system directories
    r'^/usr',
    r'^/bin',
    r'^/sbin',
    r'^/etc',
    r'^/var(?!/tmp)',  # Allow /var/tmp but block other /var paths
    r'^/root',
    r'^/boot',
    r'^/sys',
    r'^/proc',
    # Project directories (prevent writing into codebase)
    r'langconfig[\\/]backend',
    r'langconfig[\\/]src',
    r'[\\/]node_modules[\\/]',
    r'[\\/]\.git[\\/]',
]


def _validate_output_path(path: str) -> tuple[bool, str, Optional[str]]:
    """
    Validate a custom output path for security and writability.

    Returns:
        (is_valid, error_message, resolved_path)
    """
    from pathlib import Path as PathLib
    import os
    import re

    if not path or not path.strip():
        return False, "Path cannot be empty", None

    # Resolve to absolute path
    try:
        resolved = PathLib(path).resolve()
        resolved_str = str(resolved)
    except Exception as e:
        return False, f"Invalid path format: {str(e)}", None

    # Check for blocked patterns
    for pattern in BLOCKED_PATH_PATTERNS:
        if re.search(pattern, resolved_str, re.IGNORECASE):
            return False, f"Access to this directory is not allowed for security reasons", None

    # Check for path traversal attempts
    if '..' in path:
        return False, "Path traversal (..) is not allowed", None

    # Check if path can be created or exists
    try:
        if resolved.exists():
            if not resolved.is_dir():
                return False, "Path exists but is not a directory", None
            # Check write permissions
            test_file = resolved / ".langconfig_write_test"
            try:
                test_file.touch()
                test_file.unlink()
            except PermissionError:
                return False, "No write permission for this directory", None
            except Exception as e:
                return False, f"Cannot write to directory: {str(e)}", None
        else:
            # Try to create parent directories
            try:
                resolved.mkdir(parents=True, exist_ok=True)
                # Clean up if we just created it
                if not any(resolved.iterdir()):
                    resolved.rmdir()
            except PermissionError:
                return False, "No permission to create this directory", None
            except Exception as e:
                return False, f"Cannot create directory: {str(e)}", None

    except Exception as e:
        return False, f"Path validation error: {str(e)}", None

    return True, "", resolved_str


@router.post("/{workflow_id}/validate-output-path", response_model=ValidatePathResponse)
async def validate_output_path(
    workflow_id: int,
    request: ValidatePathRequest,
    db: Session = Depends(get_db)
):
    """
    Validate a custom output path for a workflow.

    Checks that the path:
    - Is not a system directory
    - Is not inside the project codebase
    - Can be created or already exists
    - Has write permissions

    Returns validation result with resolved absolute path.
    """
    from pathlib import Path as PathLib

    # Verify workflow exists
    workflow = db.query(WorkflowProfile).filter(
        WorkflowProfile.id == workflow_id
    ).first()

    if not workflow:
        raise HTTPException(status_code=404, detail="Workflow not found")

    # Validate the path
    is_valid, error, resolved_path = _validate_output_path(request.path)

    if is_valid:
        resolved = PathLib(resolved_path)
        return ValidatePathResponse(
            valid=True,
            resolved_path=resolved_path,
            error=None,
            writable=True,
            exists=resolved.exists()
        )
    else:
        return ValidatePathResponse(
            valid=False,
            resolved_path=None,
            error=error,
            writable=False,
            exists=False
        )


# =============================================================================
# Workflow Versioning Endpoints
# =============================================================================

# Pydantic Schemas for Versioning
class WorkflowVersionCreate(BaseModel):
    config_snapshot: dict = Field(..., description="Complete workflow configuration snapshot")
    notes: Optional[str] = Field(None, description="User notes about this version")
    created_by: Optional[str] = Field(None, description="User ID or name")


class WorkflowVersionResponse(BaseModel):
    model_config = ConfigDict(from_attributes=True)

    id: int
    workflow_id: int
    version_number: int
    config_snapshot: dict
    notes: Optional[str]
    changelog: Optional[str]
    is_current: bool
    created_by: Optional[str]
    created_at: datetime


class WorkflowExecutionCreate(BaseModel):
    version_id: int
    execution_results: dict
    token_usage: Optional[dict] = None
    cost: Optional[float] = None
    execution_time: Optional[float] = None
    status: str = "success"
    error_message: Optional[str] = None
    started_at: Optional[datetime] = None

    @field_validator('token_usage')
    @classmethod
    def validate_token_usage(cls, v):
        """
        Validate token usage structure and values.

        Metrics data validation
        - Ensures token_usage has correct structure
        - Validates token counts are non-negative integers
        - Checks total_tokens matches sum of input + output
        """
        if v is None:
            return v

        if not isinstance(v, dict):
            raise ValueError("token_usage must be a dictionary")

        # Expected keys
        valid_keys = {"input_tokens", "output_tokens", "total_tokens"}
        provided_keys = set(v.keys())

        # Check for unexpected keys
        unexpected = provided_keys - valid_keys
        if unexpected:
            logger.warning(f"Unexpected keys in token_usage: {unexpected}")

        # Validate individual token counts
        for key in ["input_tokens", "output_tokens", "total_tokens"]:
            if key in v:
                value = v[key]
                if not isinstance(value, (int, float)):
                    raise ValueError(f"{key} must be a number, got {type(value).__name__}")
                if value < 0:
                    raise ValueError(f"{key} cannot be negative, got {value}")
                # Convert to int if it's a whole number
                if isinstance(value, float) and value.is_integer():
                    v[key] = int(value)

        # Validate total matches sum (if all fields present)
        if all(k in v for k in ["input_tokens", "output_tokens", "total_tokens"]):
            expected_total = v["input_tokens"] + v["output_tokens"]
            actual_total = v["total_tokens"]
            if abs(expected_total - actual_total) > 0.01:  # Allow small floating point errors
                logger.warning(
                    f"Token usage mismatch: expected total={expected_total}, "
                    f"got total={actual_total}. Auto-correcting."
                )
                # Auto-correct the total
                v["total_tokens"] = expected_total

        return v

    @field_validator('cost')
    @classmethod
    def validate_cost(cls, v):
        """
        Validate cost is non-negative.

        Metrics data validation
        - Ensures cost is a non-negative number
        - Prevents invalid cost data
        """
        if v is None:
            return v

        if not isinstance(v, (int, float)):
            raise ValueError(f"cost must be a number, got {type(v).__name__}")

        if v < 0:
            raise ValueError(f"cost cannot be negative, got {v}")

        # Reasonable upper bound check (costs > $100 per execution seem suspicious)
        if v > 100.0:
            logger.warning(f"Unusually high execution cost: ${v:.2f}")

        return v

    @field_validator('execution_time')
    @classmethod
    def validate_execution_time(cls, v):
        """
        Validate execution time is non-negative.

        Metrics data validation
        - Ensures execution_time is a non-negative number
        - Prevents invalid timing data
        """
        if v is None:
            return v

        if not isinstance(v, (int, float)):
            raise ValueError(f"execution_time must be a number, got {type(v).__name__}")

        if v < 0:
            raise ValueError(f"execution_time cannot be negative, got {v}")

        # Reasonable upper bound check (executions > 1 hour seem suspicious)
        if v > 3600.0:
            logger.warning(f"Unusually long execution time: {v:.2f}s ({v/60:.2f}min)")

        return v


class WorkflowExecutionResponse(BaseModel):
    model_config = ConfigDict(from_attributes=True)

    id: int
    workflow_id: int
    version_id: int
    execution_results: dict
    token_usage: Optional[dict]
    cost: Optional[float]
    execution_time: Optional[float]
    status: str
    error_message: Optional[str]
    started_at: Optional[datetime]
    completed_at: datetime


class VersionComparisonResponse(BaseModel):
    version1: WorkflowVersionResponse
    version2: WorkflowVersionResponse
    diff: dict  # Structured diff between the two versions


class AgentCostBreakdown(BaseModel):
    name: str
    cost: float
    tokens: int


class ToolUsageBreakdown(BaseModel):
    name: str
    count: int


class ExecutionHistoryEntry(BaseModel):
    timestamp: str
    cost: float
    tokens: int


class WorkflowCostMetrics(BaseModel):
    workflow_id: int
    workflow_name: str
    totalCost: float
    totalTokens: int
    promptTokens: int
    completionTokens: int
    executionCount: int
    agents: List[AgentCostBreakdown]
    tools: List[ToolUsageBreakdown]
    executionHistory: List[ExecutionHistoryEntry]
    period_days: int = 30  # Default period for metrics


# Versioning Endpoints

@router.post("/{workflow_id}/versions", response_model=WorkflowVersionResponse, status_code=201)
async def create_workflow_version(
    workflow_id: int,
    version_data: WorkflowVersionCreate,
    db: Session = Depends(get_db)
):
    """
    Create a new version of a workflow.

    Automatically increments version number and marks this as the current version.
    Uses row-level locking to prevent race conditions in concurrent version creation.

    Raises:
        HTTPException 404: Workflow not found
        HTTPException 409: Version conflict (concurrent creation detected)
        HTTPException 500: Server error
    """
    try:
        with managed_transaction(db, f"create_version_{workflow_id}") as tx:
            # Check if workflow exists AND lock the row to prevent concurrent version creation
            workflow = db.query(WorkflowProfile).filter(
                WorkflowProfile.id == workflow_id
            ).with_for_update().first()

            if not workflow:
                raise HTTPException(status_code=404, detail="Workflow not found")

            # CRITICAL: Get the next version number with row-level lock
            # This prevents concurrent requests from reading the same version number
            last_version = db.query(WorkflowVersion).filter(
                WorkflowVersion.workflow_id == workflow_id
            ).order_by(desc(WorkflowVersion.version_number)).with_for_update().first()

            next_version_number = (last_version.version_number + 1) if last_version else 1

            # Unmark all other versions as current (still within lock)
            db.query(WorkflowVersion).filter(
                WorkflowVersion.workflow_id == workflow_id,
                WorkflowVersion.is_current == True
            ).update({"is_current": False}, synchronize_session=False)

            # Create new version
            new_version = WorkflowVersion(
                workflow_id=workflow_id,
                version_number=next_version_number,
                config_snapshot=version_data.config_snapshot,
                notes=version_data.notes,
                created_by=version_data.created_by,
                is_current=True
            )

            db.add(new_version)

        # Transaction commits here and releases locks
        db.refresh(new_version)

        logger.info(
            f"Created version {next_version_number} for workflow {workflow_id}",
            extra={
                "workflow_id": workflow_id,
                "version_number": next_version_number,
                "version_id": new_version.id
            }
        )

        return new_version

    except HTTPException:
        # Re-raise HTTP exceptions (they're intentional)
        raise

    except IntegrityError as e:
        # Unique constraint violation - concurrent version creation detected
        logger.error(
            f"Version conflict for workflow {workflow_id}: {e}",
            extra={"workflow_id": workflow_id}
        )
        raise HTTPException(
            status_code=409,
            detail="Version conflict detected. Another version was created concurrently. Please retry."
        )

    except Exception as e:
        # Unexpected error
        logger.error(
            f"Failed to create version for workflow {workflow_id}: {e}",
            exc_info=True,
            extra={"workflow_id": workflow_id}
        )
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create workflow version: {str(e)}"
        )


@router.get("/{workflow_id}/versions", response_model=List[WorkflowVersionResponse])
async def list_workflow_versions(
    workflow_id: int,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db)
):
    """
    List all versions of a workflow.

    Returns versions in reverse chronological order (newest first).
    """
    # Check if workflow exists
    workflow = db.query(WorkflowProfile).filter(WorkflowProfile.id == workflow_id).first()
    if not workflow:
        raise HTTPException(status_code=404, detail="Workflow not found")

    versions = db.query(WorkflowVersion).filter(
        WorkflowVersion.workflow_id == workflow_id
    ).order_by(desc(WorkflowVersion.version_number)).offset(skip).limit(limit).all()

    return versions


@router.get("/{workflow_id}/versions/{version_number}", response_model=WorkflowVersionResponse)
async def get_workflow_version(
    workflow_id: int,
    version_number: int,
    db: Session = Depends(get_db)
):
    """Get a specific version of a workflow."""
    version = db.query(WorkflowVersion).filter(
        WorkflowVersion.workflow_id == workflow_id,
        WorkflowVersion.version_number == version_number
    ).first()

    if not version:
        raise HTTPException(status_code=404, detail="Version not found")

    return version


@router.get("/{workflow_id}/versions/{v1}/compare/{v2}", response_model=VersionComparisonResponse)
async def compare_workflow_versions(
    workflow_id: int,
    v1: int,
    v2: int,
    db: Session = Depends(get_db)
):
    """
    Compare two versions of a workflow.

    Returns both versions and a structured diff highlighting changes.
    """
    version1 = db.query(WorkflowVersion).filter(
        WorkflowVersion.workflow_id == workflow_id,
        WorkflowVersion.version_number == v1
    ).first()

    version2 = db.query(WorkflowVersion).filter(
        WorkflowVersion.workflow_id == workflow_id,
        WorkflowVersion.version_number == v2
    ).first()

    if not version1 or not version2:
        raise HTTPException(status_code=404, detail="One or both versions not found")

    # Generate diff (simple approach - can be enhanced with proper diff algorithm)
    diff = _generate_config_diff(version1.config_snapshot, version2.config_snapshot)

    return {
        "version1": version1,
        "version2": version2,
        "diff": diff
    }


@router.post("/{workflow_id}/executions", response_model=WorkflowExecutionResponse, status_code=201)
async def create_workflow_execution(
    workflow_id: int,
    execution_data: WorkflowExecutionCreate,
    db: Session = Depends(get_db)
):
    """
    Record a workflow execution tied to a specific version.

    This tracks execution results, performance metrics, and costs.

    Raises:
        HTTPException 404: Workflow or version not found
        HTTPException 400: Version does not belong to workflow
        HTTPException 500: Server error
    """
    try:
        with managed_transaction(db, f"create_execution_{workflow_id}"):
            # Check if workflow exists
            workflow = db.query(WorkflowProfile).filter(
                WorkflowProfile.id == workflow_id
            ).first()

            if not workflow:
                raise HTTPException(status_code=404, detail="Workflow not found")

            # CRITICAL FIX: Check version exists AND belongs to this workflow
            # This prevents creating executions with versions from different workflows
            version = db.query(WorkflowVersion).filter(
                WorkflowVersion.id == execution_data.version_id,
                WorkflowVersion.workflow_id == workflow_id  # ← ADDED VALIDATION
            ).first()

            if not version:
                raise HTTPException(
                    status_code=404,
                    detail=f"Version {execution_data.version_id} not found for workflow {workflow_id}"
                )

            # Create execution record (now guaranteed to have valid foreign keys)
            execution = WorkflowExecution(
                workflow_id=workflow_id,
                version_id=execution_data.version_id,
                execution_results=execution_data.execution_results,
                token_usage=execution_data.token_usage,
                cost=execution_data.cost,
                execution_time=execution_data.execution_time,
                status=execution_data.status,
                error_message=execution_data.error_message,
                started_at=execution_data.started_at
            )

            db.add(execution)

        # Transaction commits here
        db.refresh(execution)

        logger.info(
            f"Created execution for workflow {workflow_id}, version {execution_data.version_id}",
            extra={
                "workflow_id": workflow_id,
                "version_id": execution_data.version_id,
                "execution_id": execution.id,
                "status": execution.status
            }
        )

        # Auto-cleanup old execution history if enabled
        from config import settings
        if settings.auto_cleanup_execution_history:
            try:
                cleanup_stats = cleanup_execution_history(db, workflow_id)
                if cleanup_stats["total_removed"] > 0:
                    logger.info(
                        f"Auto-cleanup removed {cleanup_stats['total_removed']} old executions for workflow {workflow_id}",
                        extra=cleanup_stats
                    )
            except Exception as cleanup_error:
                # Don't fail the request if cleanup fails
                logger.warning(
                    f"Auto-cleanup failed for workflow {workflow_id}: {cleanup_error}",
                    exc_info=True
                )

        return execution

    except HTTPException:
        # Re-raise HTTP exceptions (they're intentional)
        raise

    except Exception as e:
        # Unexpected error
        logger.error(
            f"Failed to create execution for workflow {workflow_id}: {e}",
            exc_info=True,
            extra={"workflow_id": workflow_id}
        )
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create execution: {str(e)}"
        )


@router.get("/{workflow_id}/executions", response_model=List[WorkflowExecutionResponse])
async def list_workflow_executions(
    workflow_id: int,
    version_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db)
):
    """
    List executions for a workflow, optionally filtered by version.

    Returns executions in reverse chronological order (newest first).
    """
    query = db.query(WorkflowExecution).filter(WorkflowExecution.workflow_id == workflow_id)

    if version_id:
        query = query.filter(WorkflowExecution.version_id == version_id)

    executions = query.order_by(desc(WorkflowExecution.completed_at)).offset(skip).limit(limit).all()

    return executions


@router.post("/{workflow_id}/executions/cleanup")
async def cleanup_workflow_executions(
    workflow_id: int,
    max_count: Optional[int] = None,
    retention_days: Optional[int] = None,
    db: Session = Depends(get_db)
):
    """
    Manually trigger execution history cleanup for a workflow.

    Configurable execution history limits
    - Allows manual cleanup of old execution records
    - Accepts optional overrides for max_count and retention_days
    - Returns statistics about what was cleaned up

    Args:
        workflow_id: ID of the workflow to clean up
        max_count: Override max executions to keep (defaults to config)
        retention_days: Override max age in days (defaults to config)

    Returns:
        Cleanup statistics with removed counts and remaining executions
    """
    # Verify workflow exists
    workflow = db.query(WorkflowProfile).filter(
        WorkflowProfile.id == workflow_id
    ).first()

    if not workflow:
        raise HTTPException(status_code=404, detail="Workflow not found")

    # Trigger cleanup
    try:
        cleanup_stats = cleanup_execution_history(
            db,
            workflow_id,
            max_count=max_count,
            retention_days=retention_days
        )

        logger.info(
            f"Manual cleanup triggered for workflow {workflow_id}",
            extra={
                "workflow_id": workflow_id,
                "stats": cleanup_stats
            }
        )

        return {
            "workflow_id": workflow_id,
            "cleanup_stats": cleanup_stats,
            "message": f"Cleaned up {cleanup_stats['total_removed']} old executions"
        }

    except Exception as e:
        logger.error(
            f"Failed to cleanup executions for workflow {workflow_id}: {e}",
            exc_info=True
        )
        raise HTTPException(
            status_code=500,
            detail=f"Failed to cleanup executions: {str(e)}"
        )


@router.get("/{workflow_id}/metrics/cost", response_model=WorkflowCostMetrics)
async def get_workflow_cost_metrics(
    workflow_id: int,
    days: int = 30,
    db: Session = Depends(get_db)
):
    """
    Get aggregated cost metrics for a workflow over a specified period.

    Returns:
    - Total cost and token usage
    - Cost breakdown by agent
    - Tool usage statistics
    - Recent execution history

    Args:
        workflow_id: ID of the workflow
        days: Number of days to include in metrics (default: 30)
    """
    # Check if workflow exists
    workflow = db.query(WorkflowProfile).filter(WorkflowProfile.id == workflow_id).first()
    if not workflow:
        raise HTTPException(status_code=404, detail="Workflow not found")

    # Calculate date threshold
    cutoff_date = datetime.utcnow() - timedelta(days=days)

    # Try to get executions (may be empty if orchestration doesn't create them)
    executions = db.query(WorkflowExecution).filter(
        WorkflowExecution.workflow_id == workflow_id,
        WorkflowExecution.completed_at >= cutoff_date
    ).order_by(desc(WorkflowExecution.completed_at)).all()

    # Initialize aggregation variables
    total_cost = 0.0
    total_tokens = 0
    prompt_tokens = 0
    completion_tokens = 0
    agent_costs: Dict[str, Dict[str, Any]] = defaultdict(lambda: {"cost": 0.0, "tokens": 0})
    tool_usage: Dict[str, int] = defaultdict(int)
    execution_history = []

    # If no workflow_executions exist, aggregate from execution_events (LLM_END events)
    # This is the primary data source for workflows using the custom tracing system
    if not executions:
        logger.info(f"No workflow_executions found for workflow {workflow_id}, aggregating from execution_events")

        # Query LLM_END events for token/cost data
        llm_events = db.query(ExecutionEvent).filter(
            ExecutionEvent.workflow_id == workflow_id,
            ExecutionEvent.event_type == "LLM_END",
            ExecutionEvent.timestamp >= cutoff_date
        ).all()

        # Query tool events for tool usage tracking
        tool_events = db.query(ExecutionEvent).filter(
            ExecutionEvent.workflow_id == workflow_id,
            ExecutionEvent.event_type == "on_tool_start",
            ExecutionEvent.timestamp >= cutoff_date
        ).all()

        # Count completed tasks as execution count
        from models.core import Task
        completed_tasks = db.query(Task).filter(
            Task.workflow_profile_id == workflow_id,
            Task.status == 'completed',
            Task.completed_at >= cutoff_date
        ).count()

        # Aggregate tokens and costs by agent
        for event in llm_events:
            event_data = event.event_data or {}
            agent_name = event_data.get("agent_label", "Unknown")
            model = event_data.get("model", "unknown")
            tokens = event_data.get("tokens_used", 0)
            prompt = event_data.get("prompt_tokens", 0)
            completion = event_data.get("completion_tokens", 0)

            # Aggregate totals
            total_tokens += tokens
            prompt_tokens += prompt
            completion_tokens += completion

            # Aggregate by agent
            if agent_name not in agent_costs:
                agent_costs[agent_name] = {"cost": 0.0, "tokens": 0, "model": model}
            agent_costs[agent_name]["tokens"] += tokens

        # Aggregate tool usage
        for event in tool_events:
            event_data = event.event_data or {}
            tool_name = event_data.get("tool_name", "unknown")
            if tool_name and tool_name != "unknown":
                tool_usage[tool_name] += 1

        # Calculate costs via the model registry (single pricing source).
        # Stored model strings may carry provider prefixes; the registry
        # resolves those (exact then longest-substring match) internally.
        from core.models.registry import model_registry

        for agent_name, data in agent_costs.items():
            rate = model_registry.get_blended_cost_per_1m(data.get("model", "unknown"), default=1.00)
            agent_cost = (data["tokens"] / 1_000_000) * rate
            data["cost"] = round(agent_cost, 4)
            total_cost += agent_cost

        total_cost = round(total_cost, 4)

        # Convert to response format
        agents_list = [
            AgentCostBreakdown(name=name, cost=data["cost"], tokens=data["tokens"])
            for name, data in agent_costs.items()
        ]
        agents_list.sort(key=lambda x: x.cost, reverse=True)

        tools_list = [
            ToolUsageBreakdown(name=tool_name, count=count)
            for tool_name, count in tool_usage.items()
        ]
        tools_list.sort(key=lambda x: x.count, reverse=True)

        logger.info(f"Aggregated from events: {len(llm_events)} LLM calls, {len(tool_events)} tool calls, {total_tokens} tokens, ${total_cost}")

        return WorkflowCostMetrics(
            workflow_id=workflow_id,
            workflow_name=workflow.name,
            totalCost=total_cost,
            totalTokens=total_tokens,
            promptTokens=prompt_tokens,
            completionTokens=completion_tokens,
            executionCount=completed_tasks,
            agents=agents_list,
            tools=tools_list,
            executionHistory=[],
            period_days=days
        )

    # Aggregate data from executions
    for execution in executions:
        # Aggregate costs and tokens
        if execution.cost:
            total_cost += execution.cost

        if execution.token_usage:
            token_data = execution.token_usage
            exec_total = token_data.get("total_tokens", 0)
            exec_prompt = token_data.get("input_tokens", 0) or token_data.get("prompt_tokens", 0)
            exec_completion = token_data.get("output_tokens", 0) or token_data.get("completion_tokens", 0)

            total_tokens += exec_total
            prompt_tokens += exec_prompt
            completion_tokens += exec_completion

        # Extract agent-level data from execution_results if available
        if execution.execution_results:
            results = execution.execution_results

            # Try to extract agent-specific costs
            # Structure may vary, check common patterns
            if "agent_outputs" in results:
                for agent_name, agent_data in results["agent_outputs"].items():
                    if isinstance(agent_data, dict):
                        agent_cost = agent_data.get("cost", 0)
                        agent_tokens = agent_data.get("tokens", 0)
                        agent_costs[agent_name]["cost"] += agent_cost
                        agent_costs[agent_name]["tokens"] += agent_tokens

            # Extract tool usage
            if "tool_calls" in results:
                for tool_call in results["tool_calls"]:
                    if isinstance(tool_call, dict):
                        tool_name = tool_call.get("tool", tool_call.get("name", "unknown"))
                        tool_usage[tool_name] += 1
            elif "tools_used" in results:
                for tool_name in results["tools_used"]:
                    tool_usage[tool_name] += 1

        # Add to execution history
        execution_history.append({
            "timestamp": execution.completed_at.isoformat() if execution.completed_at else datetime.utcnow().isoformat(),
            "cost": execution.cost or 0.0,
            "tokens": (execution.token_usage or {}).get("total_tokens", 0)
        })

    # Convert agent costs dict to list
    agents_list = [
        AgentCostBreakdown(
            name=agent_name,
            cost=data["cost"],
            tokens=data["tokens"]
        )
        for agent_name, data in agent_costs.items()
    ]

    # Convert tool usage dict to list
    tools_list = [
        ToolUsageBreakdown(name=tool_name, count=count)
        for tool_name, count in tool_usage.items()
    ]

    # Sort by cost/count descending
    agents_list.sort(key=lambda x: x.cost, reverse=True)
    tools_list.sort(key=lambda x: x.count, reverse=True)

    return WorkflowCostMetrics(
        workflow_id=workflow_id,
        workflow_name=workflow.name,
        totalCost=total_cost,
        totalTokens=total_tokens,
        promptTokens=prompt_tokens,
        completionTokens=completion_tokens,
        executionCount=len(executions),
        agents=agents_list,
        tools=tools_list,
        executionHistory=execution_history[:10],  # Limit to 10 most recent
        period_days=days
    )


# Helper Functions

def _generate_config_diff(config1: dict, config2: dict) -> dict:
    """
    Generate a comprehensive nested diff between two workflow configurations.

    Uses DeepDiff library for accurate detection of changes at any depth,
    including nested nodes, edges, and configuration changes.

    Args:
        config1: First configuration (usually older version)
        config2: Second configuration (usually newer version)

    Returns:
        dict: Structured diff with added, removed, modified, type_changes, and summary

    Example:
        Old: {"nodes": [{"id": "1", "name": "Agent 1", "config": {"model": "gpt-4"}}]}
        New: {"nodes": [{"id": "1", "name": "Agent 1", "config": {"model": "gpt-5.4"}}]}
        Diff detects: nodes[0].config.model changed from "gpt-4" to "gpt-5.4"
    """
    try:
        from deepdiff import DeepDiff

        # Generate deep diff with comprehensive settings
        deep_diff = DeepDiff(
            config1,
            config2,
            ignore_order=False,  # Order matters for nodes/edges
            report_repetition=False,
            verbose_level=2,  # Detailed diff information
            view='tree'  # Tree view for better navigation
        )

        # Convert DeepDiff output to our API format (backward compatible)
        diff = {
            "added": {},
            "removed": {},
            "modified": {},
            "type_changes": {},
            "summary": {}
        }

        # Extract added items
        if 'dictionary_item_added' in deep_diff:
            for item in deep_diff['dictionary_item_added']:
                path = str(item.path(output_format='list'))
                diff["added"][path] = item.t2 if hasattr(item, 't2') else None

        # Extract removed items
        if 'dictionary_item_removed' in deep_diff:
            for item in deep_diff['dictionary_item_removed']:
                path = str(item.path(output_format='list'))
                diff["removed"][path] = item.t1 if hasattr(item, 't1') else None

        # Extract modified values
        if 'values_changed' in deep_diff:
            for item in deep_diff['values_changed']:
                path = str(item.path(output_format='list'))
                diff["modified"][path] = {
                    "old": item.t1 if hasattr(item, 't1') else None,
                    "new": item.t2 if hasattr(item, 't2') else None
                }

        # Extract type changes
        if 'type_changes' in deep_diff:
            for item in deep_diff['type_changes']:
                path = str(item.path(output_format='list'))
                diff["type_changes"][path] = {
                    "old_type": str(type(item.t1).__name__) if hasattr(item, 't1') else None,
                    "new_type": str(type(item.t2).__name__) if hasattr(item, 't2') else None,
                    "old_value": item.t1 if hasattr(item, 't1') else None,
                    "new_value": item.t2 if hasattr(item, 't2') else None
                }

        # Add summary statistics
        diff["summary"] = {
            "total_changes": (
                len(diff["added"]) +
                len(diff["removed"]) +
                len(diff["modified"]) +
                len(diff["type_changes"])
            ),
            "items_added": len(diff["added"]),
            "items_removed": len(diff["removed"]),
            "items_modified": len(diff["modified"]),
            "type_changes": len(diff["type_changes"]),
            "has_changes": bool(
                diff["added"] or
                diff["removed"] or
                diff["modified"] or
                diff["type_changes"]
            )
        }

        logger.info(
            f"Generated deep diff: {diff['summary']['total_changes']} changes detected",
            extra={"summary": diff["summary"]}
        )

        return diff

    except ImportError:
        # Fallback to simple diff if deepdiff not installed
        logger.warning("deepdiff not installed, using simple diff. Install with: pip install deepdiff")

        diff = {
            "added": {},
            "removed": {},
            "modified": {},
            "unchanged_keys": []
        }

        all_keys = set(config1.keys()) | set(config2.keys())

        for key in all_keys:
            if key not in config1:
                diff["added"][key] = config2[key]
            elif key not in config2:
                diff["removed"][key] = config1[key]
            elif config1[key] != config2[key]:
                diff["modified"][key] = {
                    "old": config1[key],
                    "new": config2[key]
                }
            else:
                diff["unchanged_keys"].append(key)

        return diff


@router.get("/executions/{execution_id}/export/docx")
async def export_execution_to_word(
    execution_id: int,
    db: Session = Depends(get_db)
):
    """
    Export a workflow execution result to a Word document (.docx)
    This endpoint uses Task IDs from the orchestration system
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from models.core import Task

        # Fetch the task (using execution_id which is actually a task_id)
        task = db.query(Task).filter(Task.id == execution_id).first()
        if not task:
            raise HTTPException(status_code=404, detail="Task not found")

        # Get the workflow
        workflow = db.query(WorkflowProfile).filter(WorkflowProfile.id == task.workflow_profile_id).first()

        # Log task data for debugging
        logger.info(f"Exporting task {execution_id} with status: {task.status}")
        logger.info(f"Task result type: {type(task.result)}")
        if task.result:
            logger.info(f"Task result keys: {task.result.keys() if isinstance(task.result, dict) else 'Not a dict'}")

        # Create document
        doc = Document()

        # Add title
        title = doc.add_heading(workflow.name if workflow else "Workflow Results", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add metadata section
        doc.add_heading('Execution Details', level=1)

        metadata_table = doc.add_table(rows=5, cols=2)
        # Use a simpler table style that's guaranteed to exist
        try:
            metadata_table.style = 'Light Grid Accent 1'
        except:
            # Fallback to basic table style
            metadata_table.style = 'Table Grid'

        # Date
        metadata_table.rows[0].cells[0].text = 'Date:'
        metadata_table.rows[0].cells[1].text = task.completed_at.strftime('%Y-%m-%d %H:%M:%S') if task.completed_at else (task.created_at.strftime('%Y-%m-%d %H:%M:%S') if task.created_at else 'N/A')

        # Status
        metadata_table.rows[1].cells[0].text = 'Status:'
        metadata_table.rows[1].cells[1].text = task.status.value if hasattr(task.status, 'value') else str(task.status)

        # Duration
        duration_text = 'N/A'
        if task.created_at and task.completed_at:
            duration_seconds = (task.completed_at - task.created_at).total_seconds()
            if duration_seconds < 60:
                duration_text = f"{round(duration_seconds)}s"
            else:
                minutes = int(duration_seconds // 60)
                seconds = int(duration_seconds % 60)
                duration_text = f"{minutes}m {seconds}s"
        metadata_table.rows[2].cells[0].text = 'Duration:'
        metadata_table.rows[2].cells[1].text = duration_text

        # Tokens
        tokens = 'N/A'
        if task.result and isinstance(task.result, dict):
            workflow_summary = task.result.get('workflow_summary', {})
            if workflow_summary.get('total_tokens'):
                tokens = f"{workflow_summary['total_tokens']:,}"
        metadata_table.rows[3].cells[0].text = 'Tokens:'
        metadata_table.rows[3].cells[1].text = tokens

        # Cost
        cost = 'N/A'
        if task.result and isinstance(task.result, dict):
            workflow_summary = task.result.get('workflow_summary', {})
            if workflow_summary.get('total_cost_usd') is not None:
                cost = f"${workflow_summary['total_cost_usd']:.4f}"
        metadata_table.rows[4].cells[0].text = 'Cost:'
        metadata_table.rows[4].cells[1].text = cost

        doc.add_paragraph()  # Add space

        # Add results section
        doc.add_heading('Results', level=1)

        # Get the formatted content
        content = ''
        if task.result and isinstance(task.result, dict):
            content = task.result.get('formatted_content', '') or task.result.get('final_output', '')
            logger.info(f"Content length: {len(content) if content else 0}")

        if not content:
            content = 'No output available'

        # Ensure content is a string
        if not isinstance(content, str):
            logger.warning(f"Content is not a string, converting: {type(content)}")
            content = str(content)

        # Parse markdown-like content and add to document
        try:
            lines = content.split('\n')
        except Exception as e:
            logger.error(f"Error splitting content: {e}")
            lines = ['Error processing content']

        for i, line in enumerate(lines):
            try:
                # Handle headers
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                # Handle images ![alt](url)
                elif re.match(r'!\[.*?\]\((.*?)\)', line):
                    match = re.match(r'!\[.*?\]\((.*?)\)', line)
                    if match:
                        image_url = match.group(1)
                        try:
                            # Download image
                            response = requests.get(image_url, timeout=10)
                            if response.status_code == 200:
                                image_stream = io.BytesIO(response.content)
                                doc.add_picture(image_stream, width=Inches(5))
                                # Add caption
                                caption = doc.add_paragraph()
                                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        except Exception as e:
                            logger.warning(f"Failed to download image {image_url}: {e}")
                            doc.add_paragraph(f"[Image: {image_url}]")
                # Handle regular text
                elif line.strip():
                    # Remove markdown formatting
                    clean_line = line
                    clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_line)  # Bold
                    clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)  # Italic
                    clean_line = re.sub(r'`(.*?)`', r'\1', clean_line)  # Code
                    doc.add_paragraph(clean_line)
                else:
                    # Empty line - add small space
                    doc.add_paragraph()
            except Exception as line_error:
                logger.warning(f"Error processing line {i}: {line_error}")
                # Continue with next line instead of failing entirely
                continue

        # Save to BytesIO
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Create filename
        filename = f"{workflow.name.replace(' ', '_')}_{execution_id}.docx" if workflow else f"workflow_results_{execution_id}.docx"

        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except HTTPException:
        # Re-raise HTTP exceptions as-is (404, etc.)
        raise
    except ImportError as e:
        logger.error(f"Import error in Word export: {e}")
        raise HTTPException(status_code=500, detail="python-docx library not installed")
    except Exception as e:
        logger.error(f"Error exporting to Word: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


# =============================================================================
# Workflow Export/Import Endpoints
# =============================================================================

class WorkflowExportOptions(BaseModel):
    """Options for workflow export."""
    include_custom_tools: bool = True
    include_metadata: bool = True


@router.post("/{workflow_id}/export/package")
async def export_workflow_package(
    workflow_id: int,
    export_mode: str = "standard",
    db: Session = Depends(get_db)
):
    """
    Export workflow as a complete, executable Python package (ZIP).

    Returns a ZIP file containing all necessary code, configuration,
    and dependencies to run the workflow standalone.

    Uses LangChain v1.1 / LangGraph v1.x / DeepAgents v2.x patterns.

    Args:
        workflow_id: ID of the workflow to export
        export_mode: Export mode - 'standard' (fixed config) or 'configurable' (runtime config UI)
    """
    from core.codegen.workflow_exporter import ExecutableWorkflowExporter

    # Get workflow
    workflow = db.query(WorkflowProfile).filter(
        WorkflowProfile.id == workflow_id
    ).first()

    if not workflow:
        raise HTTPException(status_code=404, detail="Workflow not found")

    # Convert to dict for exporter
    workflow_dict = {
        "id": workflow.id,
        "name": workflow.name,
        "description": workflow.description,
        "strategy_type": workflow.strategy_type.value if hasattr(workflow.strategy_type, "value") else str(workflow.strategy_type),
        "configuration": workflow.configuration or {},
        "blueprint": workflow.blueprint or {},
    }

    # DEBUG: Log what's being passed to the exporter
    config = workflow.configuration or {}
    nodes = config.get("nodes", [])
    logger.info(f"[EXPORT DEBUG] Workflow {workflow_id} has {len(nodes)} nodes in configuration")
    logger.info(f"[EXPORT DEBUG] Export mode: {export_mode}")
    for i, node in enumerate(nodes[:3]):  # First 3 nodes
        node_config = node.get("config", {})
        logger.info(f"[EXPORT DEBUG] Node {i} ({node.get('id')}): config.model={node_config.get('model')}")
        logger.info(f"[EXPORT DEBUG] Node {i} ({node.get('id')}): config.system_prompt={node_config.get('system_prompt', '')[:50]}...")

    try:
        # Generate export with specified mode
        exporter = ExecutableWorkflowExporter(
            workflow=workflow_dict,
            project_id=workflow.project_id or 0,
            export_mode=export_mode
        )

        zip_bytes = await exporter.export_to_zip()

        # Return as downloadable file
        safe_name = workflow.name.lower().replace(" ", "_").replace("-", "_")
        filename = f"workflow_{safe_name}_{workflow_id}.zip"

        return StreamingResponse(
            io.BytesIO(zip_bytes),
            media_type="application/zip",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )

    except Exception as e:
        logger.error(f"Failed to export workflow {workflow_id}: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Export failed: {str(e)}"
        )


@router.get("/{workflow_id}/export/config")
async def export_workflow_config(
    workflow_id: int,
    include_custom_tools: bool = True,
    include_metadata: bool = True,
    db: Session = Depends(get_db)
):
    """
    Export workflow as .langconfig JSON for sharing between LangConfig instances.

    This format preserves the complete workflow configuration including:
    - Nodes and edges
    - Custom tool definitions
    - DeepAgent configurations

    The exported file can be imported into another LangConfig instance.
    """
    from services.workflow_config_service import WorkflowConfigService
    from fastapi.responses import JSONResponse

    try:
        service = WorkflowConfigService(db)
        config = await service.export_workflow_config(
            workflow_id=workflow_id,
            include_custom_tools=include_custom_tools,
            include_metadata=include_metadata
        )

        # Get workflow name for filename
        workflow = db.query(WorkflowProfile).filter(
            WorkflowProfile.id == workflow_id
        ).first()

        safe_name = workflow.name.lower().replace(" ", "_").replace("-", "_") if workflow else "workflow"
        filename = f"{safe_name}.langconfig"

        return JSONResponse(
            content=config,
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )

    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Failed to export workflow config {workflow_id}: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Export failed: {str(e)}"
        )


class WorkflowImportRequest(BaseModel):
    """Request body for workflow import."""
    config: dict = Field(..., description="The .langconfig JSON content")
    project_id: int = Field(..., description="Project to import into")
    name_override: Optional[str] = Field(None, description="Optional name override")
    create_custom_tools: bool = Field(True, description="Create custom tools from config")


@router.post("/import")
async def import_workflow_config(
    import_request: WorkflowImportRequest,
    db: Session = Depends(get_db)
):
    """
    Import a .langconfig file to create a new workflow.

    Validates version compatibility and creates:
    - The workflow with all nodes and edges
    - Any custom tools included in the config
    """
    from services.workflow_config_service import WorkflowConfigService

    try:
        service = WorkflowConfigService(db)
        result = await service.import_workflow_config(
            config=import_request.config,
            project_id=import_request.project_id,
            owner_id=0,  # TODO: Get from auth context
            name_override=import_request.name_override,
            create_custom_tools=import_request.create_custom_tools
        )

        return result

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Failed to import workflow: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Import failed: {str(e)}"
        )


@router.get("/{workflow_id}/export/config/preview")
async def preview_workflow_config(
    workflow_id: int,
    db: Session = Depends(get_db)
):
    """
    Preview what would be exported without downloading.

    Returns summary information about the export.
    """
    from services.workflow_config_service import WorkflowConfigService

    try:
        service = WorkflowConfigService(db)
        config = await service.export_workflow_config(
            workflow_id=workflow_id,
            include_custom_tools=True,
            include_metadata=True
        )

        # Return just the summary info
        return service.get_config_info(config)

    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Failed to preview workflow config {workflow_id}: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Preview failed: {str(e)}"
        )
